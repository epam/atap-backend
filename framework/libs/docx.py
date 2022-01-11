import docx
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_BREAK
from docx.text.run import InlineShape
from docx.oxml.shape import CT_Inline, CT_NonVisualDrawingProps, CT_Picture
from docx.oxml.xmlchemy import RequiredAttribute
from docx.oxml.simpletypes import XsdString, ST_DrawingElementId
from docx.oxml.__init__ import parse_xml, register_element_cls


class My_CT_NonVisualDrawingProps(CT_NonVisualDrawingProps):
    """
    Used for ``<wp:docPr>`` element, and perhaps others. Specifies the id and
    name of a DrawingML drawing.
    """
    id = RequiredAttribute('id', ST_DrawingElementId)
    name = RequiredAttribute('name', XsdString)
    descr = RequiredAttribute("descr", XsdString)


register_element_cls('pic:cNvPr',     My_CT_NonVisualDrawingProps)
register_element_cls('wp:docPr',      My_CT_NonVisualDrawingProps)


class My_CT_Inline(CT_Inline):
    @classmethod
    def my_new(cls, cx, cy, shape_id, pic, descr, make_as_decorative):
        """
        Return a new ``<wp:inline>`` element populated with the values passed
        as parameters.
        """
        if descr is not None:
            inline = parse_xml(cls._inline_xml_with_descr())
        elif make_as_decorative:
            inline = parse_xml(cls._inline_xml_make_as_decorative())
        else:
            inline = parse_xml(cls._inline_xml())

        inline.extent.cx = cx
        inline.extent.cy = cy
        inline.docPr.id = shape_id
        inline.docPr.name = 'Picture %d' % shape_id
        if descr is not None:
            inline.docPr.descr = descr
        inline.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        inline.graphic.graphicData._insert_pic(pic)
        return inline

    @classmethod
    def my_new_pic_inline(cls, shape_id, rId, filename, cx, cy, descr, make_as_decorative):
        """
        Return a new `wp:inline` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        inline = cls.my_new(cx, cy, shape_id, pic, descr, make_as_decorative)
        inline.graphic.graphicData._insert_pic(pic)
        return inline

    @classmethod
    def _inline_xml_with_descr(cls):
        return (
            '<wp:inline %s>\n'
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:docPr id="666" name="unnamed" descr="none"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:inline>' % nsdecls('wp', 'a', 'pic', 'r')
        )

    @classmethod
    def _inline_xml_make_as_decorative(cls):
        return (
            '<wp:inline %s>\n'
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:docPr id="666" name="unnamed">\n'
            '    <a:extLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">\n'
            '      <a:ext uri="{C183D7F6-B498-43B3-948B-1728B52AA6E4}">\n'
            '        <adec:decorative xmlns:adec="http://schemas.microsoft.com/office/drawing/2017/decorative" val="1"/>\n'
            '      </a:ext>\n'
            '    </a:extLst>\n'
            '  </wp:docPr>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:inline>' % nsdecls('wp', 'a', 'pic', 'r')
        )


def new_pic_inline(part, image_descriptor, width, height, descr, make_as_decorative):
    """Return a newly-created `w:inline` element.

    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename
    return My_CT_Inline.my_new_pic_inline(shape_id, rId, filename, cx, cy, descr, make_as_decorative)


def add_picture(run, image_path_or_stream, width=None, height=None, descr=None, make_as_decorative=False):
    inline = new_pic_inline(run.part, image_path_or_stream, width, height, descr, make_as_decorative)
    run._r.add_drawing(inline)
    return InlineShape(inline)


def list_number(document, paragraph, prev=None, level=None, num=True):
    """

    Makes a paragraph into a list item with a specific level and optional restart.

    :param document: (docx.document.Document) document to add list into
    :param paragraph: (docx.paragraph.Paragraph) paragraph to turn into a list item
    :param prev: (docx.paragraph.Paragraph/None) previous paragraph in the list
    :param level: (int/None) level of the paragraph within the outline
    :param num: (bool) numbered or bulleted
    :return: None
    """

    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable 'par.style.name'
        """
        return ('w:abstractNum[{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]]/@w:abstractNumId'
                ).format(style=paragraph.style.style_id, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable 'num'
        """
        type = 'decimal' if num else 'bullet'
        return ('w:abstractNum[{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]]/@w:abstractNumId'
                ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if prev is None or prev._p.pPr is None or prev._p.pPr.numPr is None or prev._p.pPr.numPr.numId is None:
        if level is None:
            level = 0
        numbering = document.part.numbering_part.numbering_definitions._numbering
        anum = get_abstract_id()
        num = numbering.add_num(anum)
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        num = prev._p.pPr.numPr.numId.val
    paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level


def add_bookmark(paragraph, bookmark_text, bookmark_name):
    """

    :param paragraph: (docx.paragraph.Paragraph) paragraph to turn into a bookmark
    :param bookmark_text: (string)
    :param bookmark_name: (string)
    :return: (None)
    """
    run = paragraph.add_run()

    tag = run._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)

    return run


def add_link_on_bookmark(p, link_to, text, color=RGBColor(0x44, 0x72, 0xC4), size=12, name='Segoe UI'):
    """

    :param p: (docx.paragraph.Paragraph) paragraph to turn into a link on bookmark
    :param link_to: (sting) name of bookmark
    :param text: (string)
    :param color: (RGBColor)
    :param size: (int) size of text
    :param name: (string) name of font
    :return: None
    """
    # create hyperlink node
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

    # set attribute for link to bookmark
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = p.add_run()
    r._r.append(hyperlink)
    r.font.name = name
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.underline = True


def contents(paragraph_for_title, paragraph):
    paragraph_for_title.add_run("CONTENTS")

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = r'TOC \o "1-3" \h \z \u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    run.add_break(WD_BREAK.PAGE)


def style(run, size=None, name=None, underline=False, bold=None, italic=False, color=None):
    """
    sets text size and styles

    :param run: proxy object wrapping <w:r> element
    :param size: (int) font size
    :param name: (str) the name of the font style
    :param underline: (bool) specify whether to highlight text in underline
    :param bold: (bool) specify whether to highlight text in bold
    :param italic: (bool) specify whether to highlight text in italic
    :param color: (RGBColor) color of text
    :return:
    """
    if size is not None:
        run.font.size = Pt(size)
    if name is not None:
        run.font.name = name
    if bold is not None:
        run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)


def add_hyperlink(paragraph, text, url, color=RGBColor(0x00, 0x00, 0xFF),
                  underline=True, bold=None, italic=None, size=None, name=None):
    """ This gets access to the document.xml.rels file and gets a new relation id value

    Parameters:
        paragraph - proxy object wrapping <w:p> element.
        text(str) - link text to be visible
        url(str) - url of the page where the link points
        color(rgb color) - color of text
        underline(bool)
        bold(bool) - bold text selection
        italic(bool) - select the text in italics
        size(int) - font size
        name(str) - the name of the font style

    Returns:
        OxmlElement hyperlink
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    r_pr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(r_pr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    if size is not None:
        r.font.size = Pt(size)
    if name is not None:
        r.font.name = name
    if bold is not None:
        r.font.bold = bold
    if italic is not None:
        r.font.italic = italic
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style
    r.font.color.rgb = color
    r.font.underline = underline

    return hyperlink
