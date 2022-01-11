from docx.document import Document
from bs4 import BeautifulSoup, NavigableString
from functools import reduce


from framework.libs.docx import style, add_hyperlink, list_number
from framework.report.lib import create_list


TAGS = ['ol', 'ul', 'strong', 'b', 'i', 'em', 'u', 'a']


def parse_html(document, text, paragraph=None, styles=None, name=None, size=None, bold=None, italic=None, underline=None,
               link_color=None, recursive=False, last_elem=False, parent_tag=None):
    def add_style(r, n=None, s=None, b=None, i=None, u=None):
        kwargs = dict()
        for x, key in zip([n, s, b, i, u], ["name", "size", "bold", "italic", "underline"]):
            if x is not None:
                kwargs[key] = x
        style(r, **kwargs)

    def find_root_children(bsoup):
        root = bsoup
        if isinstance(root, NavigableString):
            return [root]
        children = list(root.children)
        while len(children) == 1 and children[0].name not in TAGS:
            root = children[0]
            if isinstance(root, NavigableString):
                return [root]
            children = list(root.children)
        return children

    if not text:
        return

    p = paragraph
    if not recursive:
        repls = ('&LT;', '&lt;'), ('&GT;', '&gt;')
        text = reduce(lambda x, kv: x.replace(*kv), repls, text)

    try:
        soup = BeautifulSoup(text, 'html.parser')
    except RecursionError:
        return

    if not recursive:
        root_children = list(filter(lambda x: isinstance(x, NavigableString) or x.text.strip() or x.name == 'br',
                                    find_root_children(soup)))
    else:
        root_children = find_root_children(soup)
    styles = document.styles if isinstance(document, Document) and styles is None else styles
    for i, child in enumerate(root_children):
        if child.name is None and child:
            p = document.add_paragraph(style="Body Text") if p is None or (i and root_children[i - 1].name in ['ul', 'ol']) else p
            if i + 1 < len(root_children) and root_children[i + 1].name in ['ul', 'ol']:
                text = child.rstrip()
            elif i and root_children[i - 1].name in ['ul', 'ol']:
                text = child.lstrip()
            else:
                text = child

            add_style(p.add_run(text), n=name, s=size, b=bold)
            continue

        if child.name == 'br' and (
                recursive or i + 1 != len(root_children)) and not last_elem and parent_tag not in ['li', 'p']:
            if p is None:
                p = document.add_paragraph(style="Body Text")
            else:
                p.add_run('\n')

        if child.name == 'strong' or child.name == 'b':
            bold = True
        if child.name == 'i' or child.name == 'em':
            italic = True
        if child.name == 'u':
            underline = True

        if child.name == 'a' and child.get('href'):
            p = document.add_paragraph(style="Body Text") if p is None or (parent_tag == 'p' and recursive) else p
            add_hyperlink(p, child.text, child.get('href'), size=size, name=name, color=link_color, bold=bold,
                          italic=italic)
            continue
        elif child.name == 'a':
            p = document.add_paragraph(style="Body Text") if p is None or (parent_tag == 'p' and recursive) else p
            add_style(p.add_run(child.text), name, size, bold, italic, underline)
            continue

        if ((bold is not None or italic is not None or underline is not None) or (child.name == 'p' and child.text)
                or child.name not in TAGS):
            if child.findChildren(recursive=True):
                if i and root_children[i - 1].name in ['ul', 'ol']:
                    p = document.add_paragraph(style="Body Text")
                else:
                    parent_tag = 'p' if child.name == 'p' and parent_tag is None else parent_tag

                if len(child.findChildren(recursive=False)) == 1 and len(find_root_children(BeautifulSoup(
                        child.__str__(), 'html.parser'))) == 1:
                    ch = child.findChildren(recursive=False)[0]
                    text = ch.__str__() if ch.__str__() and ch.name != 'br' else child.text
                else:
                    text = child.__str__()
                    if child.name == 'p' and i + 1 < len(root_children) and root_children[i + 1].name in ['ol', 'ul']:
                        children = child.findChildren(recursive=False)
                        if len(children) == 1 and children[0].name == 'br':
                            text = child.text

                parse_html(document, text, p, styles, name, size, bold, italic, underline, link_color,
                           recursive=True, last_elem=i + 1 == len(root_children), parent_tag=parent_tag)
                parent_tag = None
            elif child.name != 'br':
                if (child.name == 'p' and i and not recursive) or p is None or (parent_tag == 'p' and recursive) or\
                        (i and root_children[i - 1].name in ['ul', 'ol']):
                    p = document.add_paragraph(style="Body Text")
                add_style(p.add_run(child.text), n=name, s=size, b=bold, i=italic, u=underline)
                bold, italic, underline = None, None, None
            continue

        if child.name == 'ul':
            for j, el in enumerate(child.children):
                if el.name is not None:
                    if styles is not None and 'List Bullet Black' in styles:
                        p = document.add_paragraph(style="List Bullet Black") if i or j or p is None else p
                    else:
                        p = document.add_paragraph() if i or j or p is None else p
                        create_list(p)
                        p.paragraph_format.space_after = 0
                        p.paragraph_format.line_spacing = 1
                    if el.findChildren(recursive=True):
                        parse_html(document, el.__str__(), p, styles, name, size, bold, italic, underline, link_color,
                                   recursive=True, parent_tag='li')
                    else:
                        add_style(p.add_run(el.text), name, size, bold, italic, underline)

        if child.name == 'ol':
            prev_p = None
            for j, el in enumerate(child.children):
                if el.name is not None:
                    p = document.add_paragraph(style="Body Text") if i or j or p is None else p
                    list_number(document, p, prev=prev_p, num=True)
                    if el.findChildren(recursive=True):
                        parse_html(document, el.__str__(), p, styles, name, size, bold, italic, underline, link_color,
                                   recursive=True, parent_tag='li')
                    else:
                        add_style(p.add_run(el.text), n=name, s=size)
                    prev_p = p
