from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Cm

import json
import datetime
import calendar
from bs4 import BeautifulSoup as bs4

from framework.libs.docx import style, add_hyperlink
from framework.report.lib import (
    set_column_width, set_repeat_table_header, create_list, indent_table, set_sell_color,
    add_dropdown_list_in_paragraph, add_page_count, add_page_number, set_row_height
)
from framework.xlsdata import cached_vpat_data
from web_interface.apps.task.models import Task
from web_interface.apps.report.models import VpatReportParams
from framework.report.parse_html import parse_html
from wcag_information.levels_and_versions import TABLE_A, TABLE_AA, TABLE_AAA, WCAG_A_2_1, WCAG_AA_2_1

MAJOR_FONT = 'Arial'
MINOR_FONT = 'Calibri'

OPTIONS = ['Supports', 'Partially Supports', 'Does Not Support', 'Not Applicable']


class VpatReport:
    def __init__(self, task: Task, report_params: VpatReportParams):
        self.doc = Document('framework/report/templates/templateVpat2.docx')
        self.doc._body.clear_content()
        self.task = task
        with open("framework/mappings.json", "r", encoding="utf-8") as f:
            self.vpat_mapping = json.load(f)
        self.vpat_data = cached_vpat_data
        self.report_params = report_params
        self.standart = bs4(self.report_params.standart, 'html.parser').get_text().split(',')
        self.product_types = report_params.product_type.split(',') if report_params.product_type else []
        self.section_508_chapters_list = task.test_results.section508chapters_set.all()
        self.success_criteria_level_list = task.test_results.successcriterialevel_set.filter(
            product_type__in=self.product_types if self.product_types else ['Web'])
        self.date = self.task.date_started if self.task.date_started is not None else datetime.datetime.now()
        self.date = f'{calendar.month_name[self.date.month]} {self.date.year}'
        if self.report_params.project:
            self.project = self.report_params.project
        else:
            self.project = self.report_params.job.project

    def __guidelines_table_wcag(self, table, version):
        links = {'2.0': 'https://www.w3.org/TR/2008/REC-WCAG20-20081211/', '2.1': 'https://www.w3.org/TR/WCAG21/'}
        row_cells = table.add_row().cells
        add_hyperlink(row_cells[0].paragraphs[0], text=f'Web Content Accessibility Guidelines {version}',
                      url=links[version])
        p = row_cells[1].paragraphs[0]
        for i in range(1, 4):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if f'WCAG {version} {i * "A"}' in self.standart:
                p.add_run(f'Level {i * "A"} (Yes)')
            else:
                p.add_run(f'Level {i * "A"} (No)')
            p = row_cells[1].add_paragraph(style='Normal') if i < 3 else None
        self.__zeroize_space(row_cells[0].paragraphs)
        self.__zeroize_space(row_cells[1].paragraphs)

    def __guidelines_table_508(self, table):
        row_cells = table.add_row().cells
        add_hyperlink(row_cells[0].paragraphs[0],
                      text='Revised Section 508 standards published January 18, 2017 and corrected January 22, 2018',
                      url='https://www.access-board.gov/guidelines-and-standards/communications-and-it/'
                          'about-the-ict-refresh/final-rule/text-of-the-standards-and-guidelines')
        if 'Revised Section 508' in self.standart:
            row_cells[1].paragraphs[0].add_run('Yes')
        else:
            row_cells[1].paragraphs[0].add_run('No')
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.__zeroize_space(row_cells[0].paragraphs)
        self.__zeroize_space(row_cells[1].paragraphs)

    def __guidelines_table_en(self, table):
        row_cells = table.add_row().cells
        add_hyperlink(row_cells[0].paragraphs[0],
                      text='EN 301 549 Accessibility requirements suitable for public procurement of ICT products and '
                           'services in Europe, - V3.1.1 (2019-11)',
                      url='https://www.etsi.org/deliver/etsi_en/301500_301599/301549/03.01.01_60/en_301549v030101p.pdf')
        if 'EN 301 549' in self.standart:
            row_cells[1].paragraphs[0].add_run('Yes')
        else:
            row_cells[1].paragraphs[0].add_run('No')
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.__zeroize_space(row_cells[0].paragraphs)
        self.__zeroize_space(row_cells[1].paragraphs)

    def __guidelines_table(self, kind):
        table = self.doc.add_table(rows=1, cols=2, style='Standard Table')
        for cell, title in zip(table.rows[0].cells, ['Standard/Guideline', 'Included In Report']):
            cell.paragraphs[0].style = 'Heading 2'
            style(cell.paragraphs[0].add_run(title), size=12)
        self.__guidelines_table_wcag(table, '2.0')
        if kind in ['INT', 'EN', 'WCAG']:
            self.__guidelines_table_wcag(table, '2.1')
        if kind in ['INT', '508']:
            self.__guidelines_table_508(table)
        if kind in ['INT', 'EN']:
            self.__guidelines_table_en(table)

        set_column_width(table, 0, Cm(13.73))
        set_column_width(table, 1, Cm(7.45))
        set_repeat_table_header(table.rows[0])

    def __terms_placeholder(self):
        style(self.doc.add_paragraph().add_run("The terms used in the Conformance Level information are defined as "
                                               "follows:"), name=MAJOR_FONT, size=12)
        terms = [("Supports:", "The functionality of the product has at least one method that meets the criterion "
                               "without known defects or meets with equivalent facilitation."),
                 ("Partially Supports:", "Some functionality of the product does not meet the criterion."),
                 ("Does Not Support:", "The majority of product functionality does not meet the criterion."),
                 ("Not Applicable:", "The criterion is not relevant to the product."),
                 ("Not Evaluated:", "The product has not been evaluated against the criterion. This can be used only "
                                    "in WCAG 2.0 Level AAA")]

        for (title, text) in terms:
            p = self.doc.add_paragraph()
            create_list(p)
            style(p.add_run(title + " "), name=MAJOR_FONT, size=12, bold=True)
            style(p.add_run(text), name=MAJOR_FONT, size=12)
            p.paragraph_format.space_after = 0
            p.paragraph_format.line_spacing = 1

    def __terms(self):
        self.doc.add_paragraph(style='Heading 2').add_run("Terms")
        if self.report_params.terms:
            parse_html(self.doc, self.report_params.terms, name=MAJOR_FONT, size=12)
        else:
            self.__terms_placeholder()

    def __report_info(self, kind):
        self.doc.add_paragraph(style='Heading 2').add_run(f"WCAG 2.{0 if kind == '508' else 'x'} Report")
        info = {"INT": {"title": "Tables 1 and 2 also document conformance with:",
                        "items": ["EN 301 549:  Chapter 9 - Web, Sections 10.1-10.4 of Chapter 10 - Non-Web documents, "
                                  "and Sections 11.1-11.4 and 11.8.2 of Chapter 11 - Non-Web Software (open and closed "
                                  "functionality), and Sections 12.1.2 and 12.2.4 of Chapter 12 – Documentation",
                                  "Revised Section 508: Chapter 5 – 501.1 Scope, 504.2 Content Creation or Editing, and"
                                  " Chapter 6 – 602.3 Electronic Support Documentation."]},
                "508": {"title": "Tables 1 and 2 also document conformance with Revised Section 508:",
                        "items": ["Chapter 5 – 501.1 Scope, 504.2 Content Creation or Editing",
                                  "Chapter 6 – 602.3 Electronic Support Documentation"]},
                "EN": {"title": "Tables 1 and 2 document conformance with EN 301 549:",
                       "items": ["Chapter 9 - Web", "Sections 10.1-10.4 of Chapter 10 - Non-Web documents",
                                 "Sections 11.1- 11.4 and 11.8.2 of Chapter 11 - Non-Web Software (open and closed "
                                 "functionality)", "Sections 12.1.2 and 12.2.4 of Chapter 12 - Documentation"]}}
        if kind in info:
            style(self.doc.add_paragraph(style='Normal').add_run(info[kind]["title"]), name=MAJOR_FONT, size=12)
            for item in info[kind]["items"]:
                p = self.doc.add_paragraph()
                create_list(p)
                style(p.add_run(item), name=MAJOR_FONT, size=12)
                p.paragraph_format.space_after = 0
                p.paragraph_format.line_spacing = 1

        p = self.doc.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(12)
        style(p.add_run(f'Note: When reporting on conformance with the WCAG 2.{0 if kind == "508" else "x"} Success '
                        f'Criteria, they are scoped for full pages, complete processes, and accessibility-supported '
                        f'ways of using technology as documented in the '), name=MAJOR_FONT, size=12)
        add_hyperlink(p, text='WCAG 2.0 Conformance Requirements.',
                      url='https://www.w3.org/TR/WCAG20/#conformance-reqs', name=MAJOR_FONT, size=12)

    def __information_page(self, kind):
        self.doc.add_section(WD_SECTION.NEW_PAGE)
        self.doc.add_paragraph(style='Heading 1').add_run(f"{self.project.company} Accessibility "
                                                          f"Conformance Report")
        title = 'International Edition' if kind == 'INT' else 'Revised Section 508 Edition' if kind == '508' else \
            'EN 301 549 Edition' if kind == 'EN' else 'WCAG Edition'
        self.doc.add_paragraph(style='Heading 1').add_run(title)

        vpat_version = "(Based on VPAT® Version 2.4)"
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style(p.add_run(vpat_version), name=MAJOR_FONT, size=12, bold=True)

        name_and_version = self.report_params.product_name_version if self.report_params.product_name_version else \
            f"{self.project.name}/{self.project.version}"
        for (title, value) in zip(
                ("Name of Product/Version:", "Report Date:", "Product Description:", "Contact Information:", "Notes:",
                 "Evaluation Methods Used:"),
                (name_and_version, self.date, self.report_params.product_description,
                 self.project.contact, self.report_params.notes, self.report_params.evaluation_methods)):
            p = self.doc.add_paragraph(style='Heading 2')
            p.add_run(title + ' ')
            parse_html(self.doc, value, p, name=MINOR_FONT, size=18, bold=False)

        self.doc.add_paragraph(style='Heading 2').add_run("Applicable Standards/Guidelines")
        self.doc.add_paragraph(style='Normal').add_run("This report covers the degree of conformance for the following"
                                                       " accessibility standard/guidelines:")
        self.__guidelines_table(kind)
        self.__terms()
        self.__report_info(kind)
        self.doc.add_page_break()

    def __create_table_title(self, table):
        for cell, title in zip(table.rows[0].cells, ['Criteria', 'Conformance Level', 'Remarks and Explanations']):
            cell.paragraphs[0].style = 'Heading 2'
            style(cell.paragraphs[0].add_run(title), size=12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def __add_items_in_criteria_cell(self, items, cell):
        if not items:
            cell.paragraphs[-1].add_run(' - Does not apply')
        for item in items:
            p = cell.add_paragraph()
            create_list(p)
            p.add_run(item)

    def __en_criteria(self, items, cell):
        cell.add_paragraph(style='Normal').add_run('EN 301 549 Criteria')
        self.__add_items_in_criteria_cell(items, cell)

    def __508_criteria(self, items, cell):
        cell.add_paragraph(style='Normal').add_run('Revised Section 508')
        self.__add_items_in_criteria_cell(items, cell)

    def __criteria(self, wcag, cell, kind):
        cell.paragraphs[0].style = 'Normal'
        add_hyperlink(cell.paragraphs[0], text=f"{wcag} {self.vpat_mapping[wcag]['Name']}",
                      url=self.vpat_mapping[wcag]["Link"], bold=True)
        level = f" (Level {self.vpat_mapping[wcag]['Level']}" + \
                (" 2.1 only" if wcag in WCAG_A_2_1 or wcag in WCAG_AA_2_1 else "") + ")"
        cell.paragraphs[0].add_run(level)
        if kind == 'WCAG':
            self.__zeroize_space(cell.paragraphs)
            return
        cell.add_paragraph(style='Normal').add_run('Also applies to:')
        with open("framework/report/info/criteria.json", "r", encoding="utf-8") as f:
            criteria_info = json.load(f)
            if kind in ['INT', 'EN']:
                self.__en_criteria(criteria_info[wcag]["EN"], cell)
            if kind in ['INT', '508']:
                self.__508_criteria(criteria_info[wcag]["508"], cell)
        self.__zeroize_space(cell.paragraphs)

    def sort_criterias(self, criterias):
        order_of_product_type = {'Web': 1, 'Electronic Docs': 2, 'Software': 3, 'Closed': 4, 'Authoring Tool': 5}
        return sorted(criterias,
                      key=lambda x: order_of_product_type[x.product_type] if x.product_type is not None else 6)

    def __conformance_level(self, criterias, cell):
        p = cell.paragraphs[0]
        for criteria in self.sort_criterias(criterias):
            p = cell.add_paragraph(style='Normal') if p is None else p
            if criteria.product_type is not None and criteria.product_type in self.product_types:
                p.add_run(criteria.product_type + ": ")
            add_dropdown_list_in_paragraph(
                p, OPTIONS, criteria.level,
                style={'highlight_color': 'yellow' if criteria.level == 'Select support level' else None}
            )
            p.style = 'Normal'
            p = None
        self.__zeroize_space(cell.paragraphs)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def __remarks(self, vpat_data, criterias, cell):
        cell.paragraphs[0].style = 'Normal'
        p = cell.paragraphs[0]
        for criteria in self.sort_criterias(criterias):
            p = cell.add_paragraph(style='Normal') if p is None else p
            if criteria.product_type is not None:
                if criteria.product_type in self.product_types:
                    p.add_run(criteria.product_type + ": ")
                if criteria.remark:
                    parse_html(cell, text=criteria.remark, paragraph=p)
                elif criteria.level in vpat_data[criteria.product_type]:
                    parse_html(cell, text=vpat_data[criteria.product_type][criteria.level], paragraph=p)
            elif criteria.remark:
                parse_html(cell, text=criteria.remark, paragraph=p)
            elif 'no product type needed' in vpat_data and criteria.level in vpat_data['no product type needed']:
                parse_html(cell, text=vpat_data['no product type needed'][criteria.level], paragraph=p)
            p = None
        self.__zeroize_space(cell.paragraphs)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def __table_with_level(self, kind, level, wcags, notes):
        self.doc.add_paragraph(style='Heading 3').add_run(f"Table {len(level)}: Success Criteria, Level {level}")
        if f'WCAG 2.0 {level}' not in self.standart and f'WCAG 2.1 {level}' not in self.standart:
            self.doc.add_paragraph(style='Normal').add_run("Notes: Not Applicable")
            return

        self.doc.add_paragraph(style='Normal').add_run(
            f"Notes: {notes if notes is not None and notes != 'None' else ''}")
        table = self.doc.add_table(rows=1, cols=3, style='Standard Table')
        self.__create_table_title(table)
        for wcag in wcags:
            scls = self.success_criteria_level_list.filter(criteria=wcag)
            if wcag in self.vpat_mapping:
                row_cells = table.add_row().cells
                self.__criteria(wcag, row_cells[0], kind)
                self.__conformance_level(scls, row_cells[1])
                self.__remarks(self.vpat_data['wcag'][wcag], scls, row_cells[2])

        set_repeat_table_header(table.rows[0])
        indent_table(table, Cm(0))
        for i, width in enumerate([11.71, 4.76, 8.99]):
            set_column_width(table, i, Cm(width))

    def __zeroize_space(self, paragraphs):
        for p in paragraphs:
            p.paragraph_format.space_after = 0
            p.paragraph_format.space_before = 0

    def __cell_style(self, row_cells):
        placeholder_text = 'Heading cell – no response required'

        for cell in row_cells[1:]:
            cell.paragraphs[0].style = 'Normal'
            cell.paragraphs[0].add_run(placeholder_text)
        for cell in row_cells:
            set_sell_color(cell, color='D0CECE')
            self.__zeroize_space(cell.paragraphs)

    def __add_chapter_title(self, number, name, url, note, sub_name=False):
        p = self.doc.add_paragraph(style='Heading 3')
        p.add_run(f"Chapter {number}: ")
        add_hyperlink(p, text=name, url=url)
        if sub_name:
            style(p.add_run(' (see '), bold=False)
            style(p.add_run('WCAG 2.x section'), italic=True, bold=False)
            style(p.add_run(')'), bold=False)
        self.doc.add_paragraph(style='Normal').add_run(f"Notes: {note if note is not None and note != 'None' else ''}")

    def __add_chapter(self, number, section_chapter, info, vpat_data):
        self.__add_chapter_title(number, info["name"], info["url"], section_chapter.note)
        table = self.doc.add_table(rows=1, cols=3, style='Standard Table')
        self.__create_table_title(table)
        product_types = self.product_types + ['', 'no product type needed']
        for item in info["items"]:
            row_cells = table.add_row().cells
            if isinstance(item, str):
                row_cells[0].paragraphs[0].add_run(item)
                try:
                    section_criteria = section_chapter.section508criteria_set.filter(product_type__in=product_types)
                    self.__conformance_level(section_criteria.filter(criteria=item), row_cells[1])
                    self.__remarks(vpat_data[number][info["name"]][item], section_criteria.filter(criteria=item),
                                   row_cells[2])
                except KeyError:
                    pass

                for cell in row_cells:
                    self.__zeroize_space(cell.paragraphs)
                    cell.paragraphs[0].style = 'Normal'

            elif "url" in item:
                add_hyperlink(row_cells[0].paragraphs[0], text=item["name"], url=item["url"], name=MINOR_FONT, size=11,
                              bold=True, italic=True)
                self.__cell_style(row_cells)
            else:
                style(row_cells[0].paragraphs[0].add_run(item["name"]), name=MINOR_FONT,
                      size=11, bold=True, italic=True)
                self.__cell_style(row_cells)

        set_repeat_table_header(table.rows[0])
        indent_table(table, Cm(0))
        for i, width in enumerate([11.7, 6.85, 6.82]):
            set_column_width(table, i, Cm(width))

    def __revised_section_508(self):
        self.doc.add_paragraph(style='Heading 2').add_run('Revised Section 508 Report')
        if 'Revised Section 508' not in self.standart:
            self.doc.add_paragraph(style='Normal').add_run("Notes: Not Applicable")
            return

        self.doc.add_paragraph(style='Normal').add_run(
            f"Notes: {self.report_params.section_508_note if self.report_params.section_508_note is not None and self.report_params.section_508_note != 'None' else ''}")
        vpat_data_508 = dict()
        for k, v in self.vpat_data['508'].items():
            vpat_data_508[k] = {v['name']: v['criteria']}
        with open("framework/report/info/chapters_508.json", "r", encoding="utf-8") as f:
            chapters_508 = json.load(f)
            for (number, info) in chapters_508.items():
                if number not in self.report_params.applicable_508.split(','):
                    self.__add_chapter_title(number, info["name"], info["url"], 'Not Applicable')
                    continue
                section_508_chapter = self.section_508_chapters_list.filter(chapter=number, report_type='508').first()
                self.__add_chapter(number, section_508_chapter, info, vpat_data_508)

    def __EN_301_549(self):
        self.doc.add_paragraph(style='Heading 2').add_run('EN 301 549 Report')
        if 'EN 301 549' not in self.standart:
            self.doc.add_paragraph(style='Normal').add_run("Notes: Not Applicable")
            return

        self.doc.add_paragraph(style='Normal').add_run(
            f"Notes: {self.report_params.section_en_note if self.report_params.section_en_note is not None and self.report_params.section_en_note != 'None' else ''}")
        vpat_data_en = dict()
        for k, v in self.vpat_data['EN'].items():
            vpat_data_en[k] = {v['name']: v['criteria']}
        with open("framework/report/info/chapters_en.json", "r", encoding="utf-8") as f:
            chapters_en = json.load(f)
            for (number, info) in chapters_en.items():
                if number not in self.report_params.applicable_en.split(','):
                    self.__add_chapter_title(number, info["name"], info["url"], 'Not Applicable')
                    continue

                section_en_chapter = self.section_508_chapters_list.filter(chapter=number, report_type='EN').first()
                if number == '9':
                    self.__add_chapter_title(number, info["name"], info["url"], section_en_chapter.note,
                                             sub_name=True)
                    continue
                self.__add_chapter(number, section_en_chapter, info, vpat_data_en)

    def __legal_disclaimer(self):
        self.doc.add_paragraph(style='Heading 2').add_run('Legal Disclaimer (Company)')
        style(self.doc.add_paragraph(style='Normal').add_run("Include your company legal disclaimer here, if needed."),
              name=MAJOR_FONT, size=12, italic=True)

    def __title_footer(self):
        style(self.doc.sections[0].footer.paragraphs[0].add_run(107 * "_"), size=12, bold=True, name=MAJOR_FONT)
        style(self.doc.sections[0].footer.add_paragraph().add_run(f"Confidential\t\t\t\t\t\tDate: {self.date}"),
              name='Segoe UI', size=10)
        self.__zeroize_space(self.doc.sections[0].footer.paragraphs)
        self.doc.sections[0].footer.paragraphs[1].paragraph_format.space_after = Pt(10)
        tab_stops = self.doc.sections[0].footer.paragraphs[1].paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(8.25), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Cm(16.51), WD_TAB_ALIGNMENT.CENTER)

    def __title_page(self, kind):
        self.doc.add_picture('framework/report/epam_logo.png', height=Cm(1.25), width=Cm(3.53))
        self.doc.add_paragraph(style='Body Text')
        self.doc.add_paragraph(style='Body Text')
        table = self.doc.add_table(rows=1, cols=1)
        title = 'International' if kind == 'INT' else 'Section 508' if kind == '508' else 'EN 301 549' \
            if kind == 'EN' else 'WCAG'
        style(table.rows[0].cells[0].paragraphs[0].add_run(title), name='Segoe UI Black', size=12, bold=True)
        style(table.rows[0].cells[0].add_paragraph(style='Title').add_run("VOLUNTARY PRODUCT ACCESSIBILITY TEMPLATE"),
              size=20, name='Segoe UI Black')

        style(table.add_row().cells[0].paragraphs[0].add_run(self.project.company), name='Segoe UI Black',
              size=14)
        self.doc.add_paragraph(style='Body Text')
        for _ in range(5):
            self.doc.add_paragraph(style='Normal')

        style(self.doc.add_paragraph(style='Normal').add_run('Disclaimer:'), name='Segoe UI', size=11, bold=True)
        style(self.doc.add_paragraph(style='Body Text').add_run(
            f'EPAM Education and Learning completed the VPAT® for {self.project.company} based on the '
            f'initial audit and the fixes verification in July 2020.'), name='Segoe UI', size=10)
        p = self.doc.add_paragraph(style='Body Text')
        style(p.add_run(
            'EPAM Education and Learning does not warrant that this document is error free and makes no warranties, '
            'expressed or implied. The information contained in this document represents the current view of the '
            f'{self.project.company} product (published at '), name='Segoe UI', size=10)
        add_hyperlink(p, text=self.project.url, url=self.project.url, name='Segoe UI',
                      size=10)
        style(p.add_run(
            f') on the date of {self.date}. EPAM Education and Learning does not guarantee the accuracy of any '
            f'information presented hereinafter this date.'), name='Segoe UI', size=10)

        set_row_height(table, [2.79, 0.8])
        set_column_width(table, 0, Cm(16.51))
        self.__title_footer()

    def __footer(self):
        """
        create specific footer for the first page in second section
        :return:
        """
        section = self.doc.sections[-1]
        section.footer.is_linked_to_previous = False

        section.different_first_page_header_footer = True
        style(section.first_page_footer.paragraphs[0].add_run(33 * "_"), size=12, bold=True, name=MAJOR_FONT)
        style(section.first_page_footer.add_paragraph().add_run(
            "“Voluntary Product Accessibility Template” and “VPAT” are \nregistered service marks of the "
            "Information Technology Industry Council (ITI)\t"), name=MINOR_FONT, size=11)
        style(section.first_page_footer.paragraphs[1].add_run("Page "), name=MINOR_FONT)
        add_page_number(section.first_page_footer.paragraphs[1].add_run(), name=MINOR_FONT, bold=True)
        style(section.first_page_footer.paragraphs[1].add_run(" of "), name=MINOR_FONT)
        add_page_count(section.first_page_footer.paragraphs[1].add_run(), name=MINOR_FONT, bold=True)
        self.__zeroize_space(section.first_page_footer.paragraphs)
        section.first_page_footer.paragraphs[1].paragraph_format.space_after = Pt(10)
        tab_stops = section.first_page_footer.paragraphs[1].paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(8.25), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Cm(16.51), WD_TAB_ALIGNMENT.CENTER)

        style(section.footer.paragraphs[-1].add_run("Page "), name=MINOR_FONT)
        add_page_number(section.footer.paragraphs[-1].add_run(), name=MINOR_FONT, bold=True)
        style(section.footer.paragraphs[-1].add_run(" of "), name=MINOR_FONT)
        add_page_count(section.footer.paragraphs[-1].add_run(), name=MINOR_FONT, bold=True)
        section.footer.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def create(self, filename):
        print("Creating vpat docx report...")

        kind = self.report_params.type
        self.__title_page(kind)
        self.__information_page(kind)
        wcags = TABLE_A + WCAG_A_2_1 if 'WCAG 2.1 A' in self.standart else TABLE_A
        self.__table_with_level(
            kind, level='A',
            wcags=sorted(wcags, key=lambda x: (int(x.split('.')[0]), int(x.split('.')[1]), int(x.split('.')[2]))),
            notes=self.report_params.wcag_a_note)

        wcags = TABLE_AA + WCAG_AA_2_1 if 'WCAG 2.1 AA' in self.standart else TABLE_AA
        self.__table_with_level(
            kind, level='AA',
            wcags=sorted(wcags, key=lambda x: (int(x.split('.')[0]), int(x.split('.')[1]), int(x.split('.')[2]))),
            notes=self.report_params.wcag_aa_note)
        self.__table_with_level(kind, level='AAA', wcags=TABLE_AAA, notes=self.report_params.wcag_aaa_note)
        if kind in ['INT', '508']:
            self.__revised_section_508()
        if kind in ['INT', 'EN']:
            self.__EN_301_549()
        self.__legal_disclaimer()
        self.__footer()

        print("Saving vpat docx report...")
        self.doc.save(f"{filename}")
        print("Done!")
        print(f"Save as {filename}")
