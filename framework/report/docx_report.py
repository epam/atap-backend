import calendar
from datetime import date
from typing import List, Optional
from collections import defaultdict
from PIL import Image
from math import inf

from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches, Cm

from framework.libs.docx import style, add_hyperlink, list_number, add_link_on_bookmark, add_bookmark, contents, \
    add_picture
from framework.report.lib import add_dropdown_list_in_paragraph
from framework.report.parse_html import parse_html
from web_interface.apps.project.models import Project
from web_interface.apps.report.models import ConformanceLevel, Issue
from web_interface.apps.issue.models import Example, PageScreenshot, ExampleScreenshot
from web_interface.apps.page.models import Page
from web_interface.apps.framework_data.models import Test, TestResults
from web_interface.apps.task.models import Task
from framework import xlsdata
from wcag_information.levels_and_versions import TABLE_A, TABLE_AA, TABLE_AAA

THEME_FONT = 'Segoe UI'
TITLE_FONT = 'Arial'
URL_COLOR = RGBColor(0x00, 0x82, 0x96)
OPTIONS = ['Supports', 'Supports with Exceptions', 'Does Not Support', 'Not Applicable']


class AuditReport:
    def __init__(self, test_results: TestResults, graphs: dict, wcag_2_dot_0_only=False, dev=False, delta_starting_test_results: Optional[TestResults]=None):
        self.dev = dev
        self.doc = Document('framework/report/templates/templateAudit.docx')
        self.doc._body.clear_content()
        self.conformance_graph = graphs["conformance"]
        self.conformance_graph_alt = graphs["conformance_alt"]
        self.prioritization_graph = graphs["prioritization"]
        self.prioritization_graph_alt = graphs["prioritization_alt"]
        self.references, self.wcag_table_info, self.test_info, self.sr_versions, self.vpat_data, _, _ = xlsdata.load_metadata()

        self.groups_WCAG = Issue.objects.filter(
            test_results=test_results, is_best_practice=False).order_by('priority', 'wcag')
        self.groups_BP = Issue.objects.filter(
            test_results=test_results,is_best_practice=True).order_by('priority', 'wcag')

        self.test_results = test_results
        self.delta_starting_test_results = delta_starting_test_results

        self.issues = list(self.groups_WCAG)
        self.issues = sorted(
            self.issues,
            key=lambda x:
            (x.priority, int(x.wcag.split(',')[0].split('.')[0]), int(x.wcag.split(',')[0].split('.')[1]),
             int(x.wcag.split(',')[0].split('.')[2]))
            if all([i.isdigit() for i in "".join(sum([d.strip().split('.') for d in x.wcag.split(',')], []))])
            else ("Minor", inf, inf, inf)
        )

        self.issues_bp = list(self.groups_BP)
        self.issues_bp = sorted(
            self.issues_bp,
            key=lambda x:
            (x.priority, int(x.wcag.split(',')[0].split('.')[0]), int(x.wcag.split(',')[0].split('.')[1]),
             int(x.wcag.split(',')[0].split('.')[2]))
            if all([i.isdigit() for i in "".join(sum([d.strip().split('.') for d in x.wcag.split(',')], []))])
            else ("Minor", inf, inf, inf)
        )

        self.wcag_2_dot_0_only_items = TABLE_A + TABLE_AA + TABLE_AAA
        if wcag_2_dot_0_only:
            self.issues = list(
                filter(lambda x: any([i in self.wcag_2_dot_0_only_items for i in x.wcag.split(', ')]), self.issues))
            self.issues_bp = list(
                filter(lambda x: any([i in self.wcag_2_dot_0_only_items for i in x.wcag.split(', ')]), self.issues_bp))

        self.page_screenshots: List[PageScreenshot] = self.test_results.pagescreenshot_set.all()
        self.project: Project = self.test_results.task.target_job.project
        self.wcag_to_bookmark_links = defaultdict(list)
        self.issue_title = defaultdict(str)
        if self.test_results.task.last_reported:
            dt = self.test_results.task.last_reported
        else:
            dt = date.today()

        self.date = {'month': calendar.month_name[dt.month],
                     'year': dt.year}
        self.wcag_version = '2.0' if wcag_2_dot_0_only else '2.1'

    def __title(self):
        self.doc.add_picture('framework/report/epam_logo.png', height=Cm(0.89), width=Cm(2.54))
        p = self.doc.add_paragraph(style='Title')
        style(p.add_run("\n\n\n\n\n\n\nACCESSIBILITY AUDIT"), size=20, name=TITLE_FONT)

        p = self.doc.add_paragraph()
        style(p.add_run(f"\n{self.project.company} \n"), size=12, name=THEME_FONT, bold=True)
        style(p.add_run(f"\n Version {self.project.version}"), size=12, name=THEME_FONT)
        self.doc.add_page_break()

        def __add_date_on_first_footer():
            section = self.doc.sections[0]
            section.first_page_footer.paragraphs[0].add_run(f"CONFIDENTIAL | Effective Date: "
                                                            f"{self.date['month']}"
                                                            f"-{self.date['year']}")

        __add_date_on_first_footer()

    def __header_footer(self):
        section = self.doc.sections[0]
        style(section.header.paragraphs[0].add_run(f"Accessibility audit"
                                                   f"\n{self.project.company} \t\t"), name="Arial", size=9)

        section.header.paragraphs[0].add_run().add_picture("framework/report/epam_logo.png",
                                                           height=Inches(0.17), width=Inches(0.5))

        style(section.footer.add_paragraph().add_run(f"CONFIDENTIAL | Effective Date: "
                                                     f"{self.date['month']}"
                                                     f"-{self.date['year']}\n"), size=9, name='Trebuchet MS')

    def __introduction_placeholder(self):
        visual_impairments = " including one team member with visual impairments" if self.project.visual_impairments else ""
        introduction = f"EPAM carried out an accessibility audit for {self.project.company} in {self.date['month']} " \
                       f"{self.date['year']}. The audit team was composed of {self.project.testers} EPAM " \
                       f"accessibility QA experts{visual_impairments}. The team used a combination of assistive " \
                       f"technologies, automated testing and manual testing for this effort and this report presents " \
                       f"the results along with recommendations for remediation of identified issues."
        self.doc.add_paragraph(style="Body Text").add_run(introduction)

    def __introduction(self):
        self.doc.add_heading("INTRODUCTION", level=1)
        self.__introduction_placeholder()

        self.doc.add_page_break()

    def __summary_general_placeholder(self):
        p = self.doc.add_paragraph(style='Body Text')
        p.add_run(f"The scope of this work was {self.project.comment}. The results are summarized below.")

        p = self.doc.add_paragraph(style='Body Text')
        p.add_run(f"EPAM found {len(self.issues)} issue types that violate accessibility norms and/"
                  f"or demonstrate how it might be difficult to use the product.")

        p = self.doc.add_paragraph(style='Body Text')
        p.add_run("For example:")

        p = self.doc.add_paragraph(style='List Bullet Black')
        style(p.add_run("Visually impaired people will find certain text difficult to read due to lack of contrast "
                        "and will not be able to access the essential information via Assistive Technologies."),
              size=12, name=THEME_FONT)

        p = self.doc.add_paragraph(style='List Bullet Black')
        style(p.add_run("People with motor disabilities will have some difficulties navigating the web page with"
                        " a keyboard only."), size=12, name=THEME_FONT)

        p = self.doc.add_paragraph(style='List Bullet Black')
        style(p.add_run("Those who have cognitive disabilities will find it challenging to understand some content "
                        "due to its structure or visual presentation."), size=12, name=THEME_FONT)

        p = self.doc.add_paragraph(style='Body Text')
        p.add_run(
            "EPAM recommendation is to update the page code. Fixing the Critical, Major and Minor issues should "
            "be the first priority, fixing issues marked as Best Practice is the second priority.")

    def __general_summary(self):
        self.doc.add_heading("SUMMARY", level=1)
        self.__summary_general_placeholder()

    def __summary(self):
        self.__general_summary()
        self.__a11y_testing_standards()
        self.__list_of_sample_pages()
        self.__testsing_type()

    def __testing_standards_placeholder(self):
        p = self.doc.add_paragraph(style='Body Text')
        p.add_run(f"The audit was performed with reference to WCAG {self.wcag_version} and Best Practice.")

        p = self.doc.add_paragraph(style='Body Text')
        style(p.add_run(f"{len(self.issues)}"), bold=True)
        p.add_run(f" problems were discovered that violate WCAG {self.wcag_version} Success Criteria.\n")
        style(p.add_run(f"{len(self.issues_bp)}"), bold=True)
        p.add_run(" problems are technically not the violation but are recommended to be fixed to enhance usability.")

        p = self.doc.add_paragraph(style='Body Text')
        p.add_run("The table in the ")
        add_link_on_bookmark(p, link_to="ACT-(web)", text="Accessibility Checklist Table (Web) section",
                             color=URL_COLOR)
        p.add_run(f" refers to WCAG {self.wcag_version} AA level.")

    def __a11y_testing_standards(self):
        self.doc.add_heading("A11Y TESTING STANDARDS", level=2)
        self.__testing_standards_placeholder()
        self.doc.add_page_break()

    def __description_list_of_sample_pages(self):
        self.doc.add_paragraph(style='Body Text').add_run(f"EPAM {self.project.comment}")

    def __list_of_sample_pages(self):
        self.doc.add_heading("LIST OF SAMPLE PAGES OR ELEMENTS", level=2)
        self.__description_list_of_sample_pages()

        if self.page_screenshots:
            prev_p = None
            need_screenshots = len(self.page_screenshots) <= 5
            for page_info in self.page_screenshots:
                p = self.doc.add_paragraph()
                list_number(self.doc, p, prev=prev_p, num=True)
                prev_p = p
                add_hyperlink(p, page_info.page.name, page_info.page.url,
                              color=URL_COLOR, name=THEME_FONT, size=12)
                if need_screenshots and page_info.screenshot:
                    add_picture(p.add_run("\n"), page_info.screenshot, width=Cm(15.75), make_as_decorative=True)
                elif need_screenshots:
                    run = p.add_run("\n Page screenshot is missing.")
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            run = self.doc.add_paragraph().add_run("\n page_screenshots is empty")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        self.doc.add_page_break()

    def __testsing_type(self):
        self.doc.add_heading("TESTING TYPE AND TECHNOLOGIES", level=2)
        self.__testsing_type_placeholder()
        self.doc.add_page_break()

    def __testsing_type_placeholder(self):
        self.doc.add_paragraph(style="Body Text").add_run("The product was tested using the tools and technologies "
                                                          "listed below.")
        self.doc.add_paragraph(style="Body Text").add_run("Accessibility testing tools:")
        tools = ["Colour Contrast Analyzer", "aXe Developer Tools for Mozilla Firefox", "Text Spacing Bookmarklet"]
        for tool in tools:
            style(self.doc.add_paragraph(style="List Bullet Black").add_run(tool), name=THEME_FONT, size=12)

        p = self.doc.add_paragraph(style="Body Text")
        p.add_run("Based on the data retrieved from the most recent survey by ")
        webaim_link = "https://webaim.org/projects/screenreadersurvey9/"
        add_hyperlink(p, "Webaim", webaim_link, size=12, name=THEME_FONT, color=URL_COLOR)
        p.add_run(" conducted among people with disabilities and accessibility specialists, the most commonly used "
                  "configurations include:")
        configurations = ["For Windows: JAWS with Chrome, NVDA with Firefox", "For macOS: VoiceOver with Safari",
                          "For Android: TalkBack with Chrome", "For iOS: VoiceOver with Safari"]
        for conf in configurations:
            style(self.doc.add_paragraph(style="List Bullet Black").add_run(conf), name=THEME_FONT, size=12)

        self.doc.add_paragraph(style="Body Text").add_run("Thus, the assistive technologies used for testing are:")
        for technology in self.sr_versions:
            style(
                self.doc.add_paragraph(style="List Bullet Black").add_run(
                    f"{technology['name']} {technology['version']} with {technology['browser']['name']} "
                    f"{technology['browser']['version']}"),
                name=THEME_FONT,
                size=12
            )

    def __prioritization_of_issues_placeholder(self, priority_style):
        prioritization = {
            "Blocker": ": the issue makes it impossible for people with disabilities to use the web site.",
            "Critical": ": the issue prevents the user from operating the essential functionality of the web site, "
                        "though the access for the whole web site remains.",
            "Major": ": the issue violates the WCAG criteria, however, it does not cause severe obstacles for the users"
                     " of assistive technologies.",
            "Minor": ": the issue is a minor violation of WCAG or relates to the Best Practice. Minor issues in the "
                     "WCAG section must be fixed in order to comply with the standard. Minor issues in the Best "
                     "Practice section are recommended to be applied to improve the user experience for those who work "
                     "with assistive technologies. "
        }
        for key, value in prioritization.items():
            p = self.doc.add_paragraph(style="List Bullet Black")
            run = p.add_run(key)
            style(run, name=THEME_FONT, size=12, bold=True)
            run.underline = priority_style[key]["underline"]
            run.font.color.rgb = priority_style[key]["color"]
            style(p.add_run(value), size=12, name=THEME_FONT)

        self.doc.add_paragraph(style="Body Text").add_run("The bar chart below illustrates the total number of bugs "
                                                          "sorted by priority and by the criteria.")

    def __prioritization_of_issues(self, priority_style: dict):
        self.doc.add_heading("PRIORITIZATION OF ISSUES", level=2)
        self.__prioritization_of_issues_placeholder(priority_style)

        if self.issues:
            if self.prioritization_graph:
                add_picture(self.doc.add_paragraph().add_run(),
                            self.prioritization_graph, descr=self.prioritization_graph_alt)
            else:
                run = self.doc.add_paragraph(style="Body Text").add_run("There should be a prioritization_graph here")
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        else:
            # all tests passed
            self.doc.add_paragraph(style="Body Text").add_run("[No issues found]")

        self.doc.add_page_break()

    @staticmethod
    def __add_row_in_summary_table(table, issues, priority_style, chapter, header):
        category_cells = []
        for i, issue in enumerate(issues):
            cells = table.add_row().cells
            category_cells.append(cells[0])
            add_link_on_bookmark(cells[1].paragraphs[0], link_to=f"{issue.id}",
                                 text=f"Issue {chapter}.{i + 1} {issue.name}",
                                 color=URL_COLOR)
            priority = issue.priority
            run = cells[2].paragraphs[0].add_run(priority)
            style(run, name=THEME_FONT, size=12, bold=True)
            run.underline = priority_style[priority]["underline"]
            run.font.color.rgb = priority_style[priority]["color"]
        if len(category_cells) > 1:
            category_cell = category_cells[0].merge(category_cells[-1])
        else:
            category_cell = category_cells[0]
        style(category_cell.paragraphs[0].add_run(header), bold=True, name=THEME_FONT, size=12)

    def __summary_table(self, priority_style):
        self.doc.add_heading("SUMMARY TABLE", level=2)
        self.doc.add_paragraph(style="Body Text").add_run("The table below shows issues found and their priority.")
        table = self.doc.add_table(1, 3, style="Table_EPAM")

        # add header row
        heading_names = ['Category', 'Summary', 'Priority']
        heading_cells = table.rows[0].cells
        for cell, name in zip(heading_cells, heading_names):
            style(cell.paragraphs[0].add_run(name), name='Segoe UI', size=12, bold=True)

        # [No issues found]
        if not self.issues:
            cells = table.add_row().cells
            style(cells[0].paragraphs[0].add_run(f"WCAG {self.wcag_version}"), name=THEME_FONT, size=12, bold=True)
            style(cells[1].paragraphs[0].add_run("[No issues found]"), name=THEME_FONT, size=12)
        else:
            # add wcag 2.x category
            self.__add_row_in_summary_table(table, self.issues, priority_style, 1, f"WCAG {self.wcag_version}")

        if not self.issues_bp:
            cells = table.add_row().cells
            style(cells[0].paragraphs[0].add_run("Best Practice"), name=THEME_FONT, size=12, bold=True)
            style(cells[1].paragraphs[0].add_run("[No issues found]"), name=THEME_FONT, size=12)
        else:
            # add best practice
            self.__add_row_in_summary_table(table, self.issues_bp, priority_style, 2, "Best Practice")

        self.__set_col_widths(table, (Cm(3.17), Cm(10.95), Cm(3.33)))
        self.doc.add_page_break()

    def __results_general(self):
        self.doc.add_paragraph(style="Body Text").add_run(
            "The results of the audit are presented below. The audit includes the description of each issue, including "
            "steps to reproduce it, examples of the source code, types of disabilities involved, and priority of "
            "impact on users."
        )

    def __conformance_wcag(self):
        self.doc.add_heading(f"CONFORMANCE WCAG {self.wcag_version} AA", level=2)
        if self.conformance_graph:
            add_picture(self.doc.add_paragraph().add_run(),
                        self.conformance_graph, descr=self.conformance_graph_alt)
        else:
            run = self.doc.add_paragraph(style="Body Text").add_run("There should be a conformance_graph here")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        self.doc.add_page_break()

    def __results(self):
        self.doc.add_heading("RESULTS", level=1)
        self.__results_general()
        self.__conformance_wcag()

        priority_style = {"Blocker": {"underline": WD_UNDERLINE.DOUBLE, "color": RGBColor(0x0, 0x0, 0x0)},
                          "Critical": {"underline": WD_UNDERLINE.THICK, "color": RGBColor(0xBA, 0x26, 0x19)},
                          "Major": {"underline": WD_UNDERLINE.DASH, "color": RGBColor(0xB6, 0x5F, 0x28)},
                          "Minor": {"underline": WD_UNDERLINE.DOT_DOT_DASH, "color": RGBColor(0x42, 0x81, 0x7A)},
                          "": {"underline": WD_UNDERLINE.SINGLE, "color": RGBColor(0x0, 0x0, 0x0)},
                          "UNKNOWN": {"underline": WD_UNDERLINE.SINGLE, "color": RGBColor(0x0, 0x0, 0x0)}}

        self.__prioritization_of_issues(priority_style)
        self.__summary_table(priority_style)

    def steps_placeholder(self, example):
        page = example.pages.first()
        paragraph = self.doc.add_paragraph(style="Body Text")
        paragraph.add_run("1) Open the ")
        if page is not None:
            add_hyperlink(paragraph, text=page.url, url=page.url, color=URL_COLOR, name=THEME_FONT)
        else:
            paragraph.add_run("[URL]")

    def __add_bookmark(self, issue, i, chapter):
        name = issue.name
        # header
        p = self.doc.add_paragraph()
        add_bookmark(paragraph=p, bookmark_text=f"Issue {chapter}.{i + 1} {name}", bookmark_name=f"{issue.id}")
        p.style = self.doc.styles['Heading 3']
        if chapter == 1:
            for w in issue.wcag.split(", "):
                self.wcag_to_bookmark_links[w].append(
                    {"link": str(issue.id), "name": f"Issue {chapter}.{i + 1} {name}"})
            self.issue_title[str(issue.id)] = f"Issue {chapter}.{i + 1} {name}"

    def __add_example(self, i, example, len_examples, delta_example=False):
        counter = f' {i + 1}' if len_examples > 1 else ''

        title_run = self.doc.add_paragraph(style="Title heading 5").add_run("Example" + counter)
        font = title_run.font

        if delta_example:
            font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        else:
            font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        if not delta_example and self.delta_starting_test_results is not None:
            font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
            try:
                similar_issue_in_delta = Issue.objects.get(test_results=self.delta_starting_test_results, err_id=example.issue.err_id)
                if Example.objects.filter(issue=similar_issue_in_delta, problematic_element_selector=example.problematic_element_selector).exists():
                    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                elif Example.objects.filter(issue=similar_issue_in_delta, problematic_element_position=example.problematic_element_position).exists():
                    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            except Issue.DoesNotExist:
                pass

        # steps
        style(self.doc.add_paragraph(style="Body Text").add_run("Steps:"), bold=True)
        if example.steps:
            parse_html(self.doc, example.steps, paragraph=self.doc.add_paragraph(style="Body Text"),
                       name=THEME_FONT, link_color=URL_COLOR)
        else:
            self.steps_placeholder(example)

        # pages
        if len(example.pages.all()) > 1:
            style(self.doc.add_paragraph(style="Body Text").add_run("Pages:"), bold=True)
            prev_paragraph = None
            page: Page
            for page in example.pages.all():
                p = self.doc.add_paragraph()
                list_number(self.doc, p, prev=prev_paragraph, num=True)
                prev_paragraph = p
                add_hyperlink(p, page.name, page.url, color=URL_COLOR, name=THEME_FONT, size=12)

        # expected result
        style(self.doc.add_paragraph(style="Body Text").add_run("Expected result:"), bold=True)
        if example.expected_result:
            parse_html(self.doc, example.expected_result, paragraph=self.doc.add_paragraph(style="Body Text"),
                       name=THEME_FONT, link_color=URL_COLOR)

        # actual result
        style(self.doc.add_paragraph(style="Body Text").add_run("Actual result:"), bold=True)
        if example.actual_result:
            parse_html(self.doc, example.actual_result, paragraph=self.doc.add_paragraph(style="Body Text"),
                       name=THEME_FONT, link_color=URL_COLOR)

        # note
        if example.note:
            p = self.doc.add_paragraph(style="Body Text")
            style(p.add_run("Note: "), bold=True)
            parse_html(self.doc, example.note, paragraph=p, name=THEME_FONT, link_color=URL_COLOR)
        self.doc.add_paragraph()

        # screenshot
        screenshots = ExampleScreenshot.objects.filter(example=example)
        if screenshots:
            for screenshot in screenshots:
                try:
                    add_picture(self.doc.add_paragraph().add_run(), screenshot.screenshot,
                                make_as_decorative=True, width=Inches(6.94))
                    self.doc.add_paragraph()
                except OSError as err:
                    p = self.doc.add_paragraph(style="Body Text")
                    p.add_run(f"Error with screenshot: {err}").font.highlight_color = WD_COLOR_INDEX.RED
                    p.add_run('\n')
        else:
            if self.test_results.task.status == Task.RUNNING:
                add_picture(self.doc.add_paragraph().add_run(), "framework/placeholder.jpg",
                            make_as_decorative=True, width=Inches(6.94))
            else:
                p = self.doc.add_paragraph(style="Body Text")
                p.add_run("The screenshot of the element is missing.").font.highlight_color = WD_COLOR_INDEX.YELLOW
                p.add_run('\n')

        # Code Snippet
        self.doc.add_paragraph(style="Title heading 5").add_run("Code Snippet" + counter)
        style(self.doc.add_paragraph(style="Body Text").add_run("HTML"), bold=True)
        style(self.doc.add_paragraph().add_run(example.code_snippet), name="Courier New", size=12)
        if example.problematic_element_selector:
            style(self.doc.add_paragraph(style="Body Text").add_run("PATH"), bold=True)
            style(self.doc.add_paragraph().add_run(example.problematic_element_selector), name="Courier New", size=12)
        self.doc.add_paragraph()

        # recommendations
        if example.recommendations:
            self.doc.add_paragraph(style="Title heading 5").add_run("Recommendation" + counter)
            parse_html(self.doc, example.recommendations, paragraph=self.doc.add_paragraph(style="Body Text"),
                       name=THEME_FONT, link_color=URL_COLOR)
            self.doc.add_paragraph()

    def __create_issue(self, issue, i, chapter, severity):
        examples = Example.objects.filter(issue=issue, severity=severity).order_by('order_in_issuegroup', 'id')
        len_examples = len(examples)
        if not len_examples or not issue:
            return

        if severity == 'FAIL':
            self.__add_bookmark(issue, i, chapter)
        else:
            self.doc.add_heading(f"Issue {chapter}.{i + 1} {issue.name}", level=3)

        # Intro
        parse_html(self.doc, issue.intro, self.doc.add_paragraph(style="Body Text"),
                   name=THEME_FONT, link_color=URL_COLOR)

        if issue.example_shows:
            self.doc.add_paragraph()
            parse_html(self.doc, issue.example_shows, name=THEME_FONT, link_color=URL_COLOR)

        for i, example in enumerate(examples):
            self.__add_example(i, example, len_examples)

        try:
            similar_issue_in_delta = Issue.objects.get(test_results=self.delta_starting_test_results,
                                                       err_id=example.issue.err_id)
            extra_examples = Example.objects.filter(issue=similar_issue_in_delta, severity=severity).order_by('order_in_issuegroup', 'id')
            for extra_example in extra_examples:
                for example in examples:
                    if example.problematic_element_selector == extra_example.problematic_element_selector or example.problematic_element_position == extra_example.problematic_element_position:
                        break
                else:
                    self.__add_example(i, example, len_examples, delta_example=True)
        except Issue.DoesNotExist:
            pass


        # type of disability
        self.doc.add_paragraph(style="Title heading 5").add_run("Type of disability affected by this issue")
        if issue.type_of_disability:
            parse_html(self.doc, issue.type_of_disability, name=THEME_FONT, link_color=URL_COLOR)
        self.doc.add_paragraph()

        # Techniques
        self.doc.add_paragraph(style="Title heading 5").add_run("Techniques")
        if self.wcag_version == '2.0':
            if issue.err_id in self.test_info:
                techniques = (self.test_info[issue.err_id]["techniques_2_0"]).split('\n')
                techniques_html = []
                for techniq in techniques:
                    techniques_html.append('<a href="' + techniq + '">' + techniq + '</a>')
                techniques_html = '\n'.join(techniques_html)
            else:
                techniques_html = issue.techniques
            parse_html(self.doc, techniques_html, name=THEME_FONT, link_color=URL_COLOR)
        elif issue.techniques:
            parse_html(self.doc, issue.techniques, name=THEME_FONT, link_color=URL_COLOR)
        self.doc.add_paragraph()

        # Recommendations
        self.doc.add_paragraph(style="Title heading 5").add_run("Recommendations")
        if issue.recommendations:
            parse_html(self.doc, issue.recommendations, name=THEME_FONT, link_color=URL_COLOR)
        elif issue.err_id in self.test_info and 'recommendations' in self.test_info[issue.err_id]:
            self.doc.add_paragraph(style='Body Text').add_run(self.test_info[issue.err_id]['recommendations'])
        self.doc.add_paragraph()

        # Reference to standards
        self.doc.add_paragraph(style="Title heading 5").add_run("Reference to standards")

        if issue.references:
            parse_html(self.doc, issue.references, name=THEME_FONT, link_color=URL_COLOR)
        else:
            for wcag in issue.wcag.split(", "):
                if wcag in self.wcag_table_info:
                    parse_html(self.doc, self.wcag_table_info[wcag]["reference"], name=THEME_FONT, link_color=URL_COLOR)
        self.doc.add_paragraph()

        # priority
        self.doc.add_paragraph(style="Title heading 5").add_run("Priority")
        if issue.priority:
            parse_html(self.doc, issue.priority, name=THEME_FONT, link_color=URL_COLOR)

    def __add_issues(self, issues, header, chapter, severity='FAIL', page_break=True):
        self.doc.add_heading(f"{chapter}. {header}", level=2)

        # [No issues found]
        if not issues:
            self.doc.add_paragraph(style="Body Text").add_run("[No issues found]")
            if page_break:
                self.doc.add_page_break()
            return None

        for i, issue in enumerate(issues):
            self.__create_issue(issue, i, chapter, severity)
            if i < len(issues) or page_break:
                self.doc.add_page_break()

    def __add_warnings(self):
        issues = self.issues + self.issues_bp
        examples_without_group = Example.objects \
            .filter(test_results=self.test_results, issue=None, severity='WARN') \
            .order_by('order_in_issuegroup', 'id')
        len_examples = len(examples_without_group)
        self.__add_issues(issues, header="WARNINGS", chapter=3, severity="WARN", page_break=not len_examples)
        if len_examples:
            self.doc.add_heading("Examples without group", level=3)
            for i, example in enumerate(examples_without_group):
                self.__add_example(i, example, len_examples)
            self.doc.add_page_break()

    def __add_errors_info(self):
        tests = Test.objects.filter(test_results=self.test_results, status__in=['ERROR', 'READY', 'NOTRUN'])
        if tests:
            self.doc.add_heading("4. ERRORS", level=2)
            for test in tests:
                self.doc.add_paragraph(style="Body Text").add_run(f"{test.name} : {test.status}")
            self.doc.add_page_break()

    def __list_of_a11y_issues(self):
        self.doc.add_heading("THE LIST OF A11Y ISSUES AND RECOMMENDATIONS", level=1)
        self.__add_issues(self.issues, header=f"WCAG {self.wcag_version}", chapter=1)
        self.__add_issues(self.issues_bp, header="BEST PRACTICE", chapter=2)
        if self.dev:
            self.__add_warnings()
            self.__add_errors_info()

    def __understanding_table(self):
        self.doc.add_heading("Understanding the table", level=2)
        self.doc.add_paragraph(style="Body Text").add_run("The table below has the following headers:")

        p = self.doc.add_paragraph(style="Body Text")
        list_number(self.doc, p, num=True)
        style(p.add_run("WCAG A or AA"), size=11, name=THEME_FONT, bold=True)
        style(p.add_run(f" – Lists the checkpoints from WCAG {self.wcag_version}."), size=12, name=THEME_FONT)
        prev_p = p

        p = self.doc.add_paragraph(style="Body Text")
        list_number(self.doc, p, prev=prev_p, num=True)
        style(p.add_run("Text of the guideline"), size=11, name=THEME_FONT, bold=True)
        style(p.add_run(" – gives the description of WCAG guidelines."), size=12, name=THEME_FONT)
        prev_p = p

        p = self.doc.add_paragraph(style="Body Text")
        list_number(self.doc, p, prev=prev_p, num=True)
        style(p.add_run("Conformance"), size=11, name=THEME_FONT, bold=True)
        style(p.add_run(" – Specifies how well the product meets each guideline:"), size=12, name=THEME_FONT)
        prev_p = p

        p = self.doc.add_paragraph(style="List Bullet Black")
        style(p.add_run("Supports"), size=12, name=THEME_FONT, bold=True)
        style(p.add_run(" – The functionality of the product has at least one method that meets the criteria without "
                        "known defects or meets with equivalent facilitation."), size=12, name=THEME_FONT)

        p = self.doc.add_paragraph(style="List Bullet Black")
        style(p.add_run("Supports with exceptions"), size=12, name=THEME_FONT, bold=True)
        style(p.add_run(" – Some functionality of the product does not meet the criteria."), size=12, name=THEME_FONT)

        p = self.doc.add_paragraph(style="List Bullet Black")
        style(p.add_run("Does not support"), size=12, name=THEME_FONT, bold=True)
        style(p.add_run(" – The majority of the functionality of the product does not meet the criteria."), size=12,
              name=THEME_FONT)

        p = self.doc.add_paragraph(style="List Bullet Black")
        style(p.add_run("Not Applicable"), size=12, name=THEME_FONT, bold=True)
        style(p.add_run(" – The criteria are not relevant to the product."), size=12, name=THEME_FONT)

        p = self.doc.add_paragraph(style="Body Text")
        list_number(self.doc, p, prev=prev_p, num=True)
        style(p.add_run("Web"), size=11, name=THEME_FONT, bold=True)
        style(p.add_run(" – gives recommendations for implementing HTML-related principles and techniques for WCAG "
                        "conformance."), size=12, name=THEME_FONT)
        prev_p = p

        p = self.doc.add_paragraph(style="Body Text")
        list_number(self.doc, p, prev=prev_p, num=True)
        style(p.add_run("Issue"), size=11, name=THEME_FONT, bold=True)
        style(p.add_run(" – contains the list of related issues."), size=12, name=THEME_FONT)

    @staticmethod
    def __set_col_widths(table, widths):
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = width

    @staticmethod
    def __set_alignment_cells(table, alignment=WD_ALIGN_VERTICAL.TOP):
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = alignment

    def __add_info(self, cells, wcag, level):
        add_hyperlink(cells[0].paragraphs[0], text=wcag, url=self.wcag_table_info[wcag][f"link_{self.wcag_version}"],
                      name=THEME_FONT, size=10, color=URL_COLOR, bold=True)
        cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        parse_html(cells[1], self.wcag_table_info[wcag]["reference"], paragraph=cells[1].paragraphs[0],
                   styles=self.doc.styles, name=THEME_FONT, size=10, link_color=URL_COLOR)

        # run = cells[2].paragraphs[0].add_run(level)
        # style(run, name=THEME_FONT, size=10)
        # if level == 'Select conformance level':
        #     run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        add_dropdown_list_in_paragraph(
            cells[2].paragraphs[0], OPTIONS, level,
            style={
                'font_name': THEME_FONT,
                'size': 10,
                'highlight_color': 'yellow' if level == 'Select conformance level' else None,
            }
        )

        if 'web' in self.wcag_table_info[wcag]:
            style(cells[3].paragraphs[0].add_run(self.wcag_table_info[wcag]["web"]), name=THEME_FONT, size=10)

    def __default_add_links_to_issues(self, cell, wcag):
        if wcag in self.wcag_to_bookmark_links:
            p = cell.paragraphs[0]
            for issue in self.wcag_to_bookmark_links[wcag]:
                add_link_on_bookmark(p, link_to=issue["link"], text=issue["name"], color=URL_COLOR, size=10)
                p = cell.add_paragraph(style="Body Text")

    def __add_links_to_issues(self, cell, remark, wcag):
        p = cell.paragraphs[0]
        if remark:
            p.add_run(remark)
            p = cell.add_paragraph(style="Body Text")

        for bookmark in self.wcag_to_bookmark_links[wcag]:
            add_link_on_bookmark(p, link_to=str(bookmark['link']), text=bookmark['name'], color=URL_COLOR, size=10)
            p = cell.add_paragraph(style="Body Text")

    def __default_conformance_table(self, table):
        if self.wcag_version == '2.0':
            wcag_items = set(self.wcag_table_info.keys()).intersection(set(self.wcag_2_dot_0_only_items))
        else:
            wcag_items = self.wcag_table_info.keys()
        wcag_items = sorted(wcag_items,
                            key=lambda x: (int(x.split('.')[0]), int(x.split('.')[1]), int(x.split('.')[2])))
        for wcag in wcag_items:
            cells = table.add_row().cells
            self.__add_info(cells, wcag=wcag, level="Supports")
            self.__default_add_links_to_issues(cells[4], wcag)

    def __add_conformance_table(self):
        self.doc.add_heading("Conformance Table", level=2)
        table = self.doc.add_table(1, 5, style="Table_EPAM")

        # add header row
        for cell, name in zip(table.rows[0].cells, ['WCAG A or AA', 'Text Of The Guideline', 'Conformance', 'Web',
                                                    'Issue']):
            style(cell.paragraphs[0].add_run(name), name='Arial', size=10, bold=True)

        if self.wcag_version == '2.0':
            conformance_table_info: List[ConformanceLevel] = self.test_results.conformancelevel_set.all().filter(
                WCAG__in=self.wcag_2_dot_0_only_items).order_by('WCAG')
        else:
            conformance_table_info: List[ConformanceLevel] = self.test_results.conformancelevel_set.all().order_by(
                'WCAG')

        conformance_table_info = [conformance_item for conformance_item in conformance_table_info if
                                  conformance_item.WCAG != "undefined"]

        if not conformance_table_info:
            self.__default_conformance_table(table)
        else:
            conformance_table_info = sorted(
                conformance_table_info,
                key=lambda x: (int(x.WCAG.split('.')[0]), int(x.WCAG.split('.')[1]), int(x.WCAG.split('.')[2]))
            )
            for info in conformance_table_info:
                cells = table.add_row().cells
                self.__add_info(cells, info.WCAG, info.level)
                self.__add_links_to_issues(cells[4], info.remark, info.WCAG)

        self.__set_col_widths(table, (Cm(1.6), Cm(6.09), Cm(2.84), Cm(4.95), Cm(3.24)))
        self.__set_alignment_cells(table)

    def __accessibility_checklist_table(self):
        # header
        p = self.doc.add_paragraph()
        add_bookmark(paragraph=p,
                     bookmark_text="ACCESSIBILITY CHECKLIST TABLE (WEB)",
                     bookmark_name="ACT-(web)")
        p.style = self.doc.styles['Heading 1']

        self.__understanding_table()
        self.__add_conformance_table()
        self.doc.add_page_break()

    def __references(self):
        self.doc.add_heading("REFERENCES", level=1)
        references = self.references[1:3] if self.wcag_version == '2.1' else self.references[-2:]
        for reference in references:
            p = self.doc.add_paragraph(style="Body Text")
            p.add_run(f"{reference['name']} ")
            add_hyperlink(p, text=reference["link"], url=reference["link"], color=URL_COLOR, name=THEME_FONT)

    def __create(self):
        # if not metadata_validation(self.references, self.wcag_table_info, self.test_info, self.sr_versions):
        #     self.doc.add_paragraph(style='Body Text').add_run('The metadata is incomplete!')
        #     return

        self.__title()
        self.__header_footer()

        paragraph_for_title = self.doc.add_paragraph(style='TOC Heading')
        contents_paragraph = self.doc.add_paragraph()

        self.__introduction()
        self.__summary()
        self.__results()
        self.__list_of_a11y_issues()
        self.__accessibility_checklist_table()
        self.__references()
        contents(paragraph_for_title, contents_paragraph)

    def __fix_picture_bug(self):
        for docPr in self.doc._part._element.findall('.//' + qn('wp:docPr')):
            docPr.set('id', str(int(docPr.get('id')) + 100000))

    def create_report(self, filename='report.docx'):
        """
        create report and save as docx file
        :param filename: (str) file name *.docx
        :return:
        """
        self.__create()
        print("Saving docx report...")
        self.__fix_picture_bug()
        self.doc.save(f"{filename}")
        print("Done!")
        print(f"Save as {filename}")
